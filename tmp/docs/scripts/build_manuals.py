from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path('/Volumes/NAS(MAC)/Data/Herd/archive-master-app')
ASSETS = ROOT / 'tmp/docs/manual_assets_v2'
OUT = ROOT / 'output/doc'
OUT.mkdir(parents=True, exist_ok=True)

TODAY = date.today().strftime('%Y-%m-%d')


def style_doc(doc: Document) -> None:
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ArchiveMaster')
    r.bold = True
    r.font.size = Pt(28)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    r2.bold = True
    r2.font.size = Pt(22)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(subtitle)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(f'Version: {TODAY}')

    doc.add_page_break()


def add_image(doc: Document, filename: str, caption: str) -> None:
    path = ASSETS / filename
    if not path.exists():
        doc.add_paragraph(f'[Imagen no encontrada: {filename}]')
        return

    doc.add_picture(str(path), width=Inches(6.7))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_index(doc: Document, items: list[str]) -> None:
    doc.add_heading('Indice', level=1)
    for item in items:
        doc.add_paragraph(item, style='List Number')


def add_permissions_matrix(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build_portal_manual() -> Path:
    doc = Document()
    style_doc(doc)

    add_cover(
        doc,
        'Manual de Usuario - Portal Operativo',
        'Roles: Recepcionista, Encargado de Oficina, Encargado de Archivo, Usuario Regular',
    )

    add_index(
        doc,
        [
            'Objetivo y alcance',
            'Flujos funcionales del portal',
            'Matriz de permisos por rol operativo',
            'Guía por rol: Recepcionista',
            'Guía por rol: Encargado de Oficina',
            'Guía por rol: Encargado de Archivo',
            'Guía por rol: Usuario Regular',
            'Troubleshooting y buenas practicas',
        ],
    )

    doc.add_heading('Objetivo y alcance', level=1)
    doc.add_paragraph(
        'Este manual explica, paso a paso, cómo operar el Portal de ArchiveMaster para todos los roles funcionales '
        'excepto Super Administrador. Incluye procesos reales del sistema, capturas de pantalla y diagramas de flujo.'
    )

    doc.add_heading('Flujos funcionales del portal', level=1)
    add_image(doc, 'flow_portal_roles.png', 'Flujo operativo por rol en el portal.')
    add_image(doc, 'flow_document_lifecycle.png', 'Ciclo de vida del documento desde carga hasta archivo fisico.')
    add_image(doc, 'flow_notifications_realtime.png', 'Flujo funcional de notificaciones en tiempo real.')

    add_permissions_matrix(
        doc,
        'Matriz de permisos por rol operativo',
        ['Acción', 'Recepcionista', 'Encargado de Oficina', 'Encargado de Archivo', 'Usuario Regular'],
        [
            ['Crear documento', 'Si', 'No', 'No', 'No'],
            ['Enviar a oficinas', 'Si', 'Si (si es responsable)', 'No', 'No'],
            ['Aprobar/Rechazar destino', 'No', 'Si', 'No', 'No'],
            ['Asignar ubicacion fisica', 'No', 'No', 'Si', 'No'],
            ['Imprimir etiqueta', 'No', 'No', 'Si', 'No'],
            ['Consultar documentos propios', 'Si', 'Si', 'Si', 'Si'],
            ['Ver notificaciones', 'Si', 'Si', 'Si', 'Si'],
        ],
    )

    doc.add_heading('Guía por rol: Recepcionista', level=1)
    doc.add_paragraph('Flujo recomendado y validado en pruebas E2E:')
    for step in [
        'Ingresar al portal con las credenciales del rol Recepcionista.',
        'Crear un nuevo documento y cargar archivos.',
        'Completar metadatos obligatorios y validar el borrador.',
        'Configurar oficinas destino y registrar la nota de envío.',
        'Confirmar en revisión final y crear el documento.',
        'Enviar a una o más oficinas destino.',
        'Monitorear estado y trazabilidad en Mis Documentos.',
    ]:
        doc.add_paragraph(step, style='List Number')

    add_image(doc, 'portal_reception_01_dashboard.png', 'Dashboard del rol Recepcionista.')
    add_image(doc, 'portal_reception_02_create_step1.png', 'Paso 1: carga inicial de archivos.')
    add_image(doc, 'portal_reception_03_uploading_state.png', 'Estado observado: el último archivo puede quedarse en "Subiendo...".')
    add_image(doc, 'portal_reception_04_upload_ready_after_reload.png', 'Después de recargar, el archivo pasa a estado "Listo".')
    add_image(doc, 'portal_reception_05_metadata_step2.png', 'Paso 2: captura de metadatos del documento.')
    add_image(doc, 'portal_reception_06_metadata_completed.png', 'Paso 2 validado: metadatos completos.')
    add_image(doc, 'portal_reception_07_configuration_step3.png', 'Paso 3: configuración de oficinas destino.')
    add_image(doc, 'portal_reception_08_configuration_completed.png', 'Paso 3 completo con destinatarios.')
    add_image(doc, 'portal_reception_09_configuration_validated.png', 'Paso 3 validado con resumen de envío.')
    add_image(doc, 'portal_reception_10_review_step4.png', 'Paso 4: revisión final antes de crear.')
    add_image(doc, 'portal_reception_11_document_created_detail.png', 'Documento creado con código y trazabilidad.')
    add_image(doc, 'portal_reception_12_distribution_form.png', 'Formulario de distribución a oficinas.')
    add_image(doc, 'portal_reception_13_distribution_sent.png', 'Distribución confirmada y registrada.')

    doc.add_heading('Guía por rol: Encargado de Oficina', level=1)
    for step in [
        'Revisar notificaciones de documentos recibidos.',
        'Abrir la bandeja de documentos recibidos.',
        'Gestionar estado: marcar recibido, revisar, responder o cerrar.',
        'Agregar observaciones y continuar el flujo.',
    ]:
        doc.add_paragraph(step, style='List Number')

    add_image(doc, 'portal_office_01_login.png', 'Inicio de sesión del portal operativo.')
    add_image(doc, 'portal_office_02_dashboard.png', 'Dashboard del rol Encargado de Oficina.')
    add_image(doc, 'portal_office_03_documents.png', 'Bandeja y listado para gestión de oficina.')
    add_image(doc, 'portal_office_04_notifications.png', 'Centro de notificaciones del rol Encargado de Oficina.')
    add_image(doc, 'portal_office_05_document_actions.png', 'Acciones disponibles sobre destinos de oficina.')

    doc.add_heading('Guía por rol: Encargado de Archivo', level=1)
    for step in [
        'Confirmar recepcion del documento en archivo.',
        'Asignar ubicacion fisica final.',
        'Marcar documento como archivado.',
        'Imprimir etiqueta y validar trazabilidad.',
    ]:
        doc.add_paragraph(step, style='List Number')

    add_image(doc, 'portal_archive_01_dashboard.png', 'Dashboard del rol Encargado de Archivo.')
    add_image(doc, 'portal_archive_02_document_archive_panel.png', 'Panel de archivo físico: ubicación e impresión de etiqueta.')

    doc.add_heading('Guía por rol: Usuario Regular', level=1)
    for step in [
        'Ingresar al portal y abrir Mis Documentos.',
        'Aplicar filtros rapidos y avanzados.',
        'Consultar estado, prioridad y trazabilidad de los documentos visibles por permisos.',
    ]:
        doc.add_paragraph(step, style='List Number')

    add_image(doc, 'portal_user_01_dashboard.png', 'Dashboard del rol Usuario Regular.')
    add_image(doc, 'portal_user_02_documents_empty.png', 'Vista sin documentos cuando no tiene asignaciones.')

    doc.add_heading('Troubleshooting y buenas practicas', level=1)
    tips = [
        'Si un archivo se queda en estado "Subiendo...", validar red y recargar la vista una sola vez.',
        'Si no aparecen notificaciones en tiempo real, validar conexión de broadcast y canal de usuario.',
        'Antes de cerrar un flujo, verificar historial de actividad y estado final.',
        'No compartir credenciales entre roles; cada rol tiene permisos acotados.',
    ]
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')

    output = OUT / 'manual_portal_archivemaster.docx'
    doc.save(output)
    return output


def build_admin_manual() -> Path:
    doc = Document()
    style_doc(doc)

    add_cover(
        doc,
        'Manual de Usuario - Panel Administrativo',
        'Roles: Administrador y Administrador de Sede (sin Super Administrador)',
    )

    add_index(
        doc,
        [
            'Objetivo del panel administrativo',
            'Flujo de administración y gobierno',
            'Matriz de permisos administrativos',
            'Operación del rol Administrador',
            'Operación del rol Administrador de Sede',
            'Guía de gobierno de datos y configuración',
            'Checklist de operación diaria',
        ],
    )

    doc.add_heading('Objetivo del panel administrativo', level=1)
    doc.add_paragraph(
        'Este manual define la operacion del panel admin para los roles autorizados, excluyendo de forma explicita '
        'al Super Administrador. Incluye gestion de usuarios, documentos, workflows y reporteria.'
    )

    doc.add_heading('Flujo de administración y gobierno', level=1)
    add_image(doc, 'flow_admin_governance.png', 'Flujo recomendado de gobierno: configuracion, operacion, reporte y mejora.')

    add_permissions_matrix(
        doc,
        'Matriz de permisos administrativos',
        ['Modulo / Accion', 'Administrador', 'Administrador de Sede'],
        [
            ['Gestionar usuarios de su alcance', 'Si', 'Si'],
            ['Gestionar documentos', 'Si', 'Si'],
            ['Configurar workflows', 'Si', 'Si'],
            ['Gestionar empresas globales', 'No*', 'No'],
            ['Reportes ejecutivos', 'Si', 'Si (alcance sede)'],
            ['Configuraciones criticas de plataforma', 'No', 'No'],
        ],
    )
    doc.add_paragraph('*Sin privilegios de Super Administrador.')

    doc.add_heading('Operación del rol Administrador', level=1)
    for step in [
        'Ingresar al panel admin y revisar indicadores iniciales.',
        'Validar usuarios y roles activos.',
        'Supervisar documentos y estados de proceso.',
        'Ajustar workflows cuando cambie la operacion.',
        'Monitorear reportes y exportables.',
    ]:
        doc.add_paragraph(step, style='List Number')

    add_image(doc, 'admin_01_dashboard.png', 'Dashboard del Administrador.')
    add_image(doc, 'admin_02_users_roles.png', 'Módulo de usuarios y roles.')
    add_image(doc, 'admin_03_documents.png', 'Módulo de documentos del panel administrativo.')

    doc.add_heading('Operación del rol Administrador de Sede', level=1)
    for step in [
        'Entrar al panel con cuenta de sede.',
        'Revisar dashboard de su alcance.',
        'Gestionar usuarios y documentos de su sede.',
        'Analizar reportes para seguimiento local.',
    ]:
        doc.add_paragraph(step, style='List Number')

    add_image(doc, 'admin_branch_01_dashboard.png', 'Dashboard del Administrador de Sede.')
    add_image(doc, 'admin_branch_02_users.png', 'Usuarios visibles para el Administrador de Sede.')
    add_image(doc, 'admin_branch_03_documents.png', 'Documentos gestionados desde sede.')
    add_image(doc, 'admin_branch_04_reports.png', 'Reportería operativa por sede.')

    doc.add_heading('Guía de gobierno de datos y configuración', level=1)
    for item in [
        'Catálogos: mantener coherencia entre estados, categorías y etiquetas.',
        'Workflows: documentar cada cambio antes de pasarlo a producción.',
        'Permisos: validar por rol y no por usuario individual cuando sea posible.',
        'Trazabilidad: no eliminar históricos operativos sin respaldo.',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Checklist de operación diaria', level=1)
    checks = [
        'Revisar notificaciones no leidas.',
        'Verificar documentos en estados bloqueados.',
        'Confirmar que reportes clave se generen sin errores.',
        'Auditar cambios de roles y accesos del dia.',
    ]
    for check in checks:
        doc.add_paragraph(check, style='List Number')

    output = OUT / 'manual_admin_archivemaster.docx'
    doc.save(output)
    return output


def main() -> None:
    portal = build_portal_manual()
    admin = build_admin_manual()
    print(f'OK: {portal}')
    print(f'OK: {admin}')


if __name__ == '__main__':
    main()
